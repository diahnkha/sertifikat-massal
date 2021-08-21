"""python PyCertGen.py"""
from docx import Document
import os
import pandas as pd
import copy
import common_utils


def print_list(lst):
    for each in lst:
        print(each)


def parser(document, v=1):
    """
    v : Verbosity
    """
    LINES = []
    for p in document.paragraphs:
        inline = p.runs
        for i in range(len(inline)):
            text = inline[i].text
            if v == 1:
                print(text)
            LINES.append(text)
    return LINES


def cleanParsed(res):
    ans = []
    for each in res:
        stripped = each.strip()
        if not stripped == "":
            ans.append(stripped)
    return ans


def replacer(document, dic):
    newdocument = copy.deepcopy(document)  # deep copy
    for p in newdocument.paragraphs:
        inline = p.runs
        for i in range(len(inline)):
            text = inline[i].text
            if text in dic.keys():
                text = text.replace(text, dic[text])
                inline[i].text = text
    return newdocument

# For various people load multiple dics


def DocxLoader(FileName):
    return Document(FileName)


def SingleSubstitution(InputDocxFilePath, OutputDocxFilePath, dic):
    document = Document(InputDocxFilePath)
    IsReplaced = replacer(document, dic)
    document.save(OutputDocxFilePath)


def CertGenEngine(InputDocxFilePath, InputExcelFilePath, OutputFolder):
    document = Document(InputDocxFilePath)
    df = pd.read_excel(InputExcelFilePath, engine="openpyxl")
    Attributes = list(df.columns)
    common_utils.create_dir(OutputFolder)
    for index, row in df.iterrows():
        dic = dict(row)
        newdocument = replacer(document, dic)
        OutputFileName = str(row[Attributes[0]]) + '.docx'
        
        # choose the output file names as first column name
        OutputFilePath = os.path.join(OutputFolder, OutputFileName)
        newdocument.save(OutputFilePath)


if __name__ == '__main__':
    ROOT_DIR = os.getcwd()
    FileName = os.path.join(ROOT_DIR, "CertTemplateSamples",
                            "Certificate_of_Appreciation.docx")
    document = Document(FileName)
    olddocument = document


